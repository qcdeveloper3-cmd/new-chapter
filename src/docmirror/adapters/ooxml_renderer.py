from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docmirror.ir.models import DocumentIR


def _paragraph_text_for_cell(lines: list[str]) -> str:
    return " ".join(part for part in lines if part).strip()


def render_document(ir: DocumentIR, output_docx: Path, rtl_default: bool) -> Path:
    output_docx.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    document.add_heading("DocMirror Scaffold Output", level=1)
    document.add_paragraph(
        "TODO: replace this renderer with geometry-preserving OOXML layout logic."
    )

    for page in ir.pages:
        document.add_paragraph(f"Page {page.page_number}")

        for block in page.blocks:
            for line in block.lines:
                paragraph = document.add_paragraph(line.text)
                if rtl_default or line.writing_direction == "rtl":
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    # TODO: set BiDi paragraph and run properties at OOXML level.

        for table_ir in page.tables:
            rows = max(table_ir.n_rows, 1)
            cols = max(table_ir.n_cols, 1)
            table = document.add_table(rows=rows, cols=cols)
            for cell in table_ir.cells:
                if 0 <= cell.row < rows and 0 <= cell.col < cols:
                    text_lines = [line.text for line in cell.lines]
                    table.cell(cell.row, cell.col).text = _paragraph_text_for_cell(text_lines)
                    # TODO: apply cell merge rules, text direction, and border styles.

        for checkbox in page.checkboxes:
            mark = "[x]" if checkbox.checked else "[ ]"
            label = checkbox.label or ""
            paragraph = document.add_paragraph(f"{mark} {label}".strip())
            if rtl_default or checkbox.writing_direction == "rtl":
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if page.images:
            document.add_paragraph("TODO: image embedding placeholder")
        if page.shapes:
            document.add_paragraph("TODO: shape drawing placeholder")

    document.save(str(output_docx))
    return output_docx
